"""author: alexshcer"""

from docx import Document
import re
import os

path = '../'
direct = os.listdir(path)

for i in direct:
    document = Document()
    document.add_heading(i, 0)
    file = open('../'+i+'/question.txt').read()
    file = re.sub(r'[^\x00-\x7F]+|\x0c',' ', file) # remove all non-XML-compatible characters
    p = document.add_paragraph(file)
    document.save('../'+i+'/question.docx')